from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("onchain-time-builder-sprint-venue-partnership-proposal.docx")

# Resolved preset: narrative_proposal.
# Named override: Solana editorial brand. It preserves the preset's geometry,
# spacing and component rhythm while replacing the default blue with the
# supplied reference site's purple/green editorial palette and Arial family.
FONT = "Arial"
MONO = "Courier New"
INK = "151020"
PURPLE = "6D45B8"
PURPLE_DARK = "42296F"
GREEN = "267A55"
AMBER = "8A5A00"
MUTED = "625A6E"
LINE = "D9D2E3"
SURFACE = "F7F5FA"
PURPLE_FILL = "EEE9F8"
GREEN_FILL = "E8F5EE"
AMBER_FILL = "FFF5DD"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, name=FONT, size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGINS):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in margins.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, attrs in edges.items():
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def paragraph_bottom_border(paragraph, color=PURPLE, size=12, space=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True, page_break_before=False):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))
    if page_break_before:
        p_pr.append(OxmlElement("w:pageBreakBefore"))


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, text, fld_char_end])
    set_run_font(run, name=MONO, size=8.5, color=MUTED)


def add_hyperlink(paragraph, text, url, color=PURPLE, underline=True):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), "18")
    r_pr.append(size_el)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def add_abstract(abstract_id, num_fmt, lvl_text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        txt = OxmlElement("w:lvlText")
        txt.set(qn("w:val"), lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.extend([start, fmt, txt, suff])
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "279")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def add_num(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    add_abstract(next_abstract, "bullet", "•", FONT)
    add_num(next_num, next_abstract)
    bullet_num_id = next_num
    add_abstract(next_abstract + 1, "decimal", "%1.")
    add_num(next_num + 1, next_abstract + 1)
    return bullet_num_id, next_num + 1


def duplicate_numbering_instance(doc, source_num_id):
    numbering = doc.part.numbering_part.element
    source = None
    for node in numbering.findall(qn("w:num")):
        if int(node.get(qn("w:numId"))) == source_num_id:
            source = node
            break
    if source is None:
        raise ValueError(f"Unknown numbering instance: {source_num_id}")
    abstract_ref = source.find(qn("w:abstractNumId"))
    abstract_id = int(abstract_ref.get(qn("w:val")))
    existing = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    new_num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return new_num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, PURPLE, 18, 10),
        ("Heading 2", 13, PURPLE, 12, 6),
        ("Heading 3", 12, PURPLE_DARK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Table Body" not in styles:
        style = styles.add_style("Table Body", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Table Body"]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(9.3)
    style.font.color.rgb = rgb(INK)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.12

    if "Table Header" not in styles:
        style = styles.add_style("Table Header", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Table Header"]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(8.8)
    style.font.bold = True
    style.font.color.rgb = rgb(WHITE)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.05

    if "Small Note" not in styles:
        style = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Small Note"]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(9)
    style.font.color.rgb = rgb(MUTED)
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ONCHAIN TIME BUILDER SPRINT  /  VENUE PARTNERSHIP PROPOSAL")
    set_run_font(r, name=MONO, size=8, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("DA NANG  ·  SEPTEMBER 2026  ·  PAGE ")
    set_run_font(r, name=MONO, size=8, color=MUTED)
    add_field(p, "PAGE")


def add_para(doc, text="", size=None, color=INK, bold=False, italic=False,
             align=None, before=0, after=8, line=1.333, font=FONT, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    r = p.add_run(text)
    set_run_font(r, name=font, size=size, color=color, bold=bold, italic=italic)
    if keep:
        set_paragraph_keep(p, keep_next=True)
    return p


def add_heading(doc, text, level=1, page_break=False):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    return p


def add_bullet(doc, text, bullet_num_id, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    apply_num(p, bullet_num_id)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, size=11, color=INK, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_numbered(doc, text, number_num_id, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    apply_num(p, number_num_id)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, size=11, color=INK, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_callout(doc, label, text, fill=PURPLE_FILL, accent=PURPLE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(
        cell,
        top={"val": "nil"}, bottom={"val": "nil"},
        start={"val": "single", "sz": "24", "color": accent},
        end={"val": "nil"},
    )
    p = cell.paragraphs[0]
    p.style = doc.styles["Table Body"]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label.upper()}  ")
    set_run_font(r, name=MONO, size=9, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.3, color=INK, bold=False)
    add_para(doc, "", size=1, after=4)
    return table


def add_data_table(doc, headers, rows, widths_dxa, alignments=None, header_fill=PURPLE_DARK, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell, top={"val": "single", "sz": "4", "color": header_fill},
                         bottom={"val": "single", "sz": "4", "color": header_fill},
                         start={"val": "single", "sz": "4", "color": header_fill},
                         end={"val": "single", "sz": "4", "color": header_fill})
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Header"]
        p.alignment = (alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(str(text))
        set_run_font(r, size=8.8, color=WHITE, bold=True)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, SURFACE)
            set_cell_borders(cell, top={"val": "single", "sz": "4", "color": LINE},
                             bottom={"val": "single", "sz": "4", "color": LINE},
                             start={"val": "single", "sz": "4", "color": LINE},
                             end={"val": "single", "sz": "4", "color": LINE})
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Body"]
            p.alignment = (alignments[col_idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_section_intro(doc, kicker, title, lede, page_break=False):
    kicker_p = add_para(doc, kicker.upper(), size=9, color=PURPLE, bold=True, font=MONO, after=4)
    if page_break:
        kicker_p.paragraph_format.page_break_before = True
    add_heading(doc, title, 1)
    add_para(doc, lede, size=11.3, color=MUTED, after=12, line=1.28)


def build_document():
    doc = Document()
    setup_page(doc)
    setup_styles(doc)
    bullet_num_id, number_num_id = configure_numbering(doc)

    core = doc.core_properties
    core.title = "Onchain Time Builder Sprint — Venue Partnership Proposal"
    core.subject = "Venue partnership proposal for a two-week online program and Da Nang build weekend"
    core.author = "Onchain Time Builder Sprint organizing team"
    core.keywords = "Solana, Colosseum, Superteam Vietnam, Da Nang, builder sprint, venue partnership"
    core.comments = "Prepared 31 August 2026"

    # Cover — proposal_centerpiece pattern.
    add_para(doc, "ALIGNED WITH THE SUPERTEAM VIETNAM · COLOSSEUM SUBTRACK EFFORT", size=9, color=PURPLE,
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=MONO, after=18)
    add_para(doc, "Onchain Time\nBuilder Sprint", size=30, color=INK, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=0.96)
    add_para(doc, "Venue Partnership Proposal", size=15, color=PURPLE, bold=False,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.0)
    add_para(doc, "Learn online. Build together. Iterate toward Colosseum.", size=11.5,
             color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20, line=1.1)
    rule = add_para(doc, "", size=1, after=14)
    paragraph_bottom_border(rule, color=PURPLE, size=18, space=8)

    metadata = doc.add_table(rows=4, cols=2)
    set_table_geometry(metadata, [4680, 4680])
    meta_rows = [
        ("LOCATION", "Da Nang, Vietnam", "PROGRAM OPENS", "4 September 2026"),
        ("IN-PERSON", "18–20 September 2026", "FORMAT", "Friday night + Saturday + Sunday morning"),
        ("PARTICIPANTS", "70 builders · 11–12 teams", "CASH BUDGET", "VND 60 million, raised separately"),
        ("VENUE ASK", "In-kind space and facilities", "OPERATIONS LEAD", "Danh"),
    ]
    for row_idx, (l1, v1, l2, v2) in enumerate(meta_rows):
        for col_idx, (label, value) in enumerate(((l1, v1), (l2, v2))):
            cell = metadata.rows[row_idx].cells[col_idx]
            set_cell_shading(cell, SURFACE if row_idx % 2 == 0 else WHITE)
            set_cell_borders(cell, top={"val": "single", "sz": "4", "color": LINE},
                             bottom={"val": "single", "sz": "4", "color": LINE},
                             start={"val": "single", "sz": "4", "color": LINE},
                             end={"val": "single", "sz": "4", "color": LINE})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(label + "\n")
            set_run_font(r, name=MONO, size=8, color=PURPLE, bold=True)
            r = p.add_run(value)
            set_run_font(r, size=10.2, color=INK, bold=True)
    add_para(doc, "", size=1, after=10)
    add_callout(
        doc,
        "Partnership request",
        "Provide an accessible, builder-ready venue for Friday evening, Saturday daytime and Sunday morning, including reliable internet, power, presentation facilities, security and cleaning.",
        fill=GREEN_FILL,
        accent=GREEN,
    )
    add_para(doc, "Prepared for prospective venue partners · Prepared by Hieu and the organizing team · 31 August 2026",
             size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.1)

    # 1. Executive summary.
    add_page_break(doc)
    add_section_intro(
        doc,
        "01 / The ask",
        "A venue partnership for a focused builder program",
        "The Onchain Time Builder Sprint is a two-week online learning and product-iteration program that culminates in a compact in-person build weekend in Da Nang. The physical event is designed for collaboration, mentoring and feedback—not introductory lectures or a local prize competition.",
    )
    add_callout(
        doc,
        "Decision requested",
        "Confirm an in-kind venue partnership for up to 80 people, with access on 18–20 September 2026 and the facilities listed in this proposal.",
    )
    add_heading(doc, "Program at a glance", 2)
    add_data_table(
        doc,
        ["Dimension", "Plan"],
        [
            ("Program model", "All learning materials online from 4 September; in-person acceleration from 18–20 September"),
            ("Audience", "Approximately 70 student and independent builders in Da Nang, supported by organizers and mentors"),
            ("Expected output", "11–12 working project submissions recorded directly on the event website"),
            ("Incentive", "No local prize pool; projects continue toward the Superteam Vietnam Colosseum subtrack with a stated total prize pool of US$10,000"),
            ("Event principle", "Building and repeated iteration matter more than a final-day award"),
            ("Cash plan", "VND 60 million for catering and operations, funded separately from the venue contribution"),
        ],
        [2400, 6960],
    )
    add_heading(doc, "Why the timing works", 2)
    add_para(doc, "The in-person weekend ends on 20 September, giving teams an eight-day continuation window before the fall Colosseum hackathon begins on 28 September. Colosseum then runs through 2 November 2026. The sprint is therefore a launchpad into a longer global build cycle—not an attempt to finish a startup in one weekend.")
    add_callout(
        doc,
        "Program logic",
        "Participants prepare online, use the venue for concentrated work and feedback, then leave with a working checkpoint and a concrete Colosseum roadmap.",
        fill=GREEN_FILL,
        accent=GREEN,
    )

    # 2. Opportunity and concept.
    add_page_break(doc)
    add_section_intro(
        doc,
        "02 / Program thesis",
        "From a compelling mechanism to focused products",
        "Time.fun popularized a provocative SocialFi mechanism: creators could issue tradable tokens linked to access to their time. The mechanism attracted attention, but it also exposed difficult product questions around target users, repeatable utility, reliable delivery and sustainable market behavior.",
    )
    add_heading(doc, "The build prompt", 2)
    add_para(doc, "The sprint asks teams to treat the time-market idea as a starting point rather than a blueprint. Each team selects a specific community or professional niche, defines a concrete exchange of value, and builds the smallest experience capable of testing whether people actually want it.")
    add_callout(
        doc,
        "Core question",
        "For whom does tokenized access to time create genuine, repeatable value—and what product makes that value safe, understandable and useful?",
        fill=PURPLE_FILL,
        accent=PURPLE,
    )
    add_heading(doc, "Design principles", 2)
    principles = [
        ("Niche before scale. ", "Teams begin with a narrow user group and a specific job to be done."),
        ("Working product before pitch polish. ", "A functioning user journey matters more than a long presentation."),
        ("Evidence before assumptions. ", "Teams seek conversations, usage signals or other concrete feedback."),
        ("Iteration before awards. ", "Every checkpoint should show what changed and why."),
        ("Continuation before closure. ", "Sunday creates the next build plan rather than declaring the work finished."),
    ]
    for lead, body in principles:
        add_bullet(doc, lead + body, bullet_num_id, bold_lead=lead)
    add_heading(doc, "Participant and ecosystem outcomes", 2)
    add_data_table(
        doc,
        ["For builders", "For the local ecosystem", "For Colosseum"],
        [
            ("Practical Solana and AI-assisted building experience", "A visible, repeatable product-building format in Da Nang", "More prepared teams with clearer product scope"),
            ("A working project page and feedback history", "Connections among university and independent builders", "Projects that can use the five-week competition effectively"),
            ("A team, prototype and continuation roadmap", "Documented projects and media for future programs", "A stronger pathway into the Superteam Vietnam subtrack"),
        ],
        [3120, 3120, 3120],
        font_size=8.8,
    )

    # 3. Program journey and website.
    add_page_break(doc)
    add_section_intro(
        doc,
        "03 / Learning platform",
        "The website is the program—not just an event page",
        "All event materials are available online two weeks before the in-person weekend. The website acts as the learning environment, the project workspace, the feedback record and the direct submission channel.",
    )
    add_heading(doc, "Participant journey", 2)
    journey = [
        ("Learn", "Understand the time-market thesis, Solana foundations, product scoping and AI-agent workflows."),
        ("Form", "Create or join a team, choose a niche and open a public or organizer-visible project workspace."),
        ("Frame", "Submit the target user, problem, proposed value exchange and smallest testable product."),
        ("Build", "Publish checkpoints, receive online feedback and arrive in person with the environment already working."),
        ("Iterate", "Use Friday, Saturday and Sunday feedback to update the product and document what changed."),
        ("Continue", "Carry the project into the Superteam Vietnam subtrack and the wider Colosseum period."),
    ]
    for lead, body in journey:
        add_numbered(doc, lead + ". " + body, number_num_id, bold_lead=lead + ". ")
    add_heading(doc, "Website capabilities", 2)
    capabilities = [
        ("Learning library. ", "Topic explainers, niche prompts, Solana fundamentals, AI-agent building guidance, product-validation material and demo preparation."),
        ("Project workspaces. ", "Team profile, target user, product hypothesis, links, screenshots and build status."),
        ("Iteration and feedback. ", "Structured checkpoints show progress and evidence while mentor and peer comments remain attached to the work."),
        ("Submission and continuation. ", "The website records Sunday's submission and demo QR code, then remains open for subtrack updates."),
    ]
    for lead, body in capabilities:
        add_bullet(doc, lead + body, bullet_num_id, bold_lead=lead)
    add_page_break(doc)
    add_section_intro(
        doc,
        "03 / Online delivery",
        "Two weeks of preparation before the venue opens",
        "The pre-event period removes basic onboarding from the physical weekend. Teams progress through clear readiness milestones while the same website holds their learning material, project state and feedback.",
    )
    add_heading(doc, "Program calendar", 2)
    add_data_table(
        doc,
        ["Date", "Mode", "Milestone"],
        [
            ("4 Sep", "Online", "Learning platform, registration and project workspaces open"),
            ("4–10 Sep", "Online", "Learn, team formation and niche selection"),
            ("11–17 Sep", "Online", "Initial project hypothesis, first build and readiness check"),
            ("18–20 Sep", "Da Nang", "Kickoff, focused build day and demo morning"),
            ("21–27 Sep", "Online", "Post-event iteration and Colosseum preparation"),
            ("28 Sep–2 Nov", "Online", "Fall Colosseum hackathon period"),
        ],
        [1500, 1500, 6360],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_callout(
        doc,
        "Readiness gate",
        "By 17 September, every participating team should have a target user, a scoped product hypothesis, a working development environment and an initial website checkpoint. This is what makes the compact in-person schedule viable.",
        fill=GREEN_FILL,
        accent=GREEN,
    )
    add_heading(doc, "Pre-event deliverables", 2)
    pre_event = [
        ("By 10 September. ", "Team formed, niche selected and project workspace opened."),
        ("By 14 September. ", "Target user and product hypothesis published with at least one source of user or market evidence."),
        ("By 17 September. ", "Development environment working, first product flow started and Friday checkpoint prepared."),
    ]
    for lead, body in pre_event:
        add_bullet(doc, lead + body, bullet_num_id, bold_lead=lead)
    add_heading(doc, "Online operating rhythm", 2)
    add_bullet(doc, "Materials remain available on demand rather than being repeated as long lectures in person.", bullet_num_id)
    add_bullet(doc, "Organizer announcements and milestone reminders are published through the platform.", bullet_num_id)
    add_bullet(doc, "Mentor and peer comments stay attached to project checkpoints so iteration is visible.", bullet_num_id)

    # 4. In-person run of show.
    add_page_break(doc)
    add_section_intro(
        doc,
        "04 / In-person program",
        "Friday kickoff. Saturday build day. Sunday demo morning.",
        "The schedule protects focused build time while avoiding overnight work, Saturday dinner and Sunday lunch. Teams may continue online after the venue closes, but evening work is optional.",
    )
    add_heading(doc, "Friday, 18 September · 18:00–21:00", 2)
    add_data_table(
        doc,
        ["Time", "Activity", "Output"],
        [
            ("17:30", "Registration and light dinner", "Participants settled before the program begins"),
            ("18:00", "Welcome, program thesis and venue partner acknowledgement", "Shared context and expectations"),
            ("18:35", "Team alignment and project-page check", "Confirmed team, niche and workspace"),
            ("19:00", "Scope the smallest weekend test", "Build goal and evidence target"),
            ("20:10", "Checkpoint 1: 60-second team updates", "Target user, problem, product and Saturday plan"),
            ("20:50", "Operations briefing and wrap", "Clear arrival and support plan for Saturday"),
            ("21:00", "Venue closes", "Teams may continue online if they choose"),
        ],
        [1200, 3600, 4560],
        font_size=8.7,
    )
    add_heading(doc, "Saturday, 19 September · 09:00–17:00", 2)
    add_data_table(
        doc,
        ["Time", "Activity", "Output"],
        [
            ("09:00", "Daily briefing and build goals", "Team plan visible on the website"),
            ("09:15", "Focused building and user validation", "Working product flow and evidence"),
            ("11:30", "Mentor office hours", "Blockers and product assumptions challenged"),
            ("12:00", "Lunch", "Protected participant break"),
            ("13:00", "Build and test", "Prototype advanced from feedback"),
            ("15:00", "Tea break", "Short reset"),
            ("15:15", "Checkpoint 2: product review", "Prototype, evidence and change log"),
            ("16:15", "Apply feedback and prepare Sunday", "Submission checklist and demo plan"),
            ("16:45", "Daily close", "Owners and next actions confirmed"),
            ("17:00", "Venue closes", "No Saturday dinner required"),
        ],
        [1200, 3600, 4560],
        font_size=8.7,
    )

    # Sunday and demo format.
    add_page_break(doc)
    add_section_intro(
        doc,
        "05 / Demo morning",
        "Twelve teams, two tracks, two hours",
        "Sunday is a product-feedback and continuation session rather than a judging ceremony. Two parallel demo tracks give every team meaningful time without extending the event into lunch.",
    )
    add_heading(doc, "Sunday, 20 September · 08:30–11:30", 2)
    add_data_table(
        doc,
        ["Time", "Activity", "Output"],
        [
            ("08:30", "Coffee, light breakfast and final fixes", "Demo-ready project"),
            ("09:15", "Website submission deadline", "Stable project page and QR code"),
            ("09:20", "Demo briefing and track assignment", "Teams and audiences in position"),
            ("09:25–10:55", "Six 15-minute rounds across two parallel tracks", "Capacity for 12 complete demos"),
            ("10:55", "Feedback capture and transition", "Website feedback and next changes recorded"),
            ("11:05", "Colosseum and subtrack continuation plan", "Owners, dates and submission pathway"),
            ("11:20", "Partner acknowledgement, community photo and close", "Event completed before lunch"),
            ("11:30", "Venue closes to participants", "No Sunday lunch required"),
        ],
        [1500, 3900, 3960],
        font_size=8.7,
    )
    add_callout(
        doc,
        "Capacity math",
        "2 tracks × 6 rounds × 1 team per round = 12 team demos. Each 15-minute slot contains approximately 7 minutes of product demonstration, 5 minutes of feedback and 3 minutes of transition.",
        fill=GREEN_FILL,
        accent=GREEN,
    )
    add_heading(doc, "Demo operating model", 2)
    demo_points = [
        ("Two separated zones. ", "Each track needs a display, power, acoustic separation and one anchor mentor to protect timing."),
        ("Website-first evidence. ", "Every demo shows a project QR code; visitors leave feedback directly on the project page."),
        ("Rotating audience. ", "Teams not presenting attend other demos and contribute peer feedback."),
        ("No local ranking. ", "The close recognizes progress and readiness rather than announcing local winners."),
    ]
    for lead, body in demo_points:
        add_bullet(doc, lead + body, bullet_num_id, bold_lead=lead)
    add_heading(doc, "Feedback lenses", 2)
    add_data_table(
        doc,
        ["Lens", "Question"],
        [
            ("User clarity", "Is the target user specific and credible?"),
            ("Value", "Does the time-based exchange solve a real problem?"),
            ("Product", "Can a user complete the essential journey?"),
            ("Evidence", "What did the team learn from real feedback or behavior?"),
            ("Iteration", "What changed across checkpoints, and why?"),
            ("Continuation", "What is the next test during the Colosseum period?"),
        ],
        [2160, 7200],
        font_size=8.5,
    )

    # 6. Venue requirements.
    add_page_break(doc)
    add_section_intro(
        doc,
        "06 / Venue partnership",
        "What we are asking the venue partner to provide",
        "The venue contribution is in kind and remains separate from the event's VND 60 million catering and operations budget. The ideal partner can provide a safe, accessible and technically reliable environment for concentrated product work.",
    )
    add_heading(doc, "Access schedule", 2)
    add_data_table(
        doc,
        ["Day", "Requested access", "Purpose"],
        [
            ("Friday, 18 Sep", "16:30–21:30", "Setup, registration, kickoff and pack-down"),
            ("Saturday, 19 Sep", "08:00–17:30", "Setup, full build day and close"),
            ("Sunday, 20 Sep", "08:00–12:00", "Coffee, submissions, two demo tracks and closing"),
        ],
        [2100, 2160, 5100],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_heading(doc, "Required facilities", 2)
    venue_requirements = [
        ("Capacity and layout. ", "Seating for up to 80 people, team tables for laptop work, a presentation zone, two separated Sunday demo zones, mentor corners and a catering/water area."),
        ("Internet. ", "Stable connectivity for approximately 100 simultaneous devices, with an identified technical contact and permission for the organizers to bring backup connectivity."),
        ("Power. ", "Safe access to sufficient outlets for every team; permission to use tested extension boards and cable covers."),
        ("Presentation. ", "At least one central screen or projector with microphone support, plus a display or projector for each Sunday demo track."),
        ("Operating environment. ", "Air conditioning, restrooms, drinking-water access, accessibility information and appropriate lighting."),
        ("Safety and access. ", "Security, cleaning, entry procedures, emergency contacts, capacity limits and clear weekend opening/closing responsibility."),
    ]
    for lead, body in venue_requirements:
        add_bullet(doc, lead + body, bullet_num_id, bold_lead=lead)
    add_heading(doc, "Division of responsibility", 2)
    add_data_table(
        doc,
        ["Venue partner", "Organizing team"],
        [
            ("Space, utilities, agreed furniture and included AV", "Program design, registration, participant communication and run of show"),
            ("Weekend access, security, cleaning and venue technical contact", "Catering procurement, event crew, mentors and participant support"),
            ("Venue safety rules and emergency procedures", "Code of conduct, dietary collection, project submissions and feedback"),
            ("Brand assets and acknowledgement requirements", "Partner visibility, photography coordination and post-event report"),
        ],
        [4680, 4680],
    )

    # 7. Partner value and operations.
    add_page_break(doc)
    add_section_intro(
        doc,
        "07 / Partnership value",
        "A visible role in Da Nang's builder pipeline",
        "The venue partner is not a background supplier. It becomes the physical home of a program that connects online learning, local technical talent and a direct continuation pathway into the Solana ecosystem.",
    )
    add_heading(doc, "Venue partner benefits", 2)
    benefits = [
        ("Official recognition. ", "Named as the Official Venue Partner across the event website, registration pages and event communications."),
        ("On-site visibility. ", "Logo placement at registration, the central presentation area and both Sunday demo tracks, subject to agreed brand guidelines."),
        ("Opening participation. ", "A short welcome during Friday's kickoff and acknowledgement during Sunday's closing."),
        ("Talent access. ", "Direct engagement with approximately 70 builders and visibility into 11–12 emerging project teams."),
        ("Content and media. ", "Access to approved event photographs, project links and a concise post-event impact report."),
        ("Ecosystem association. ", "Association with a Superteam Vietnam Colosseum build effort focused on learning, iteration and real products."),
    ]
    for lead, body in benefits:
        add_bullet(doc, lead + body, bullet_num_id, bold_lead=lead)
    add_heading(doc, "Core operating team", 2)
    add_data_table(
        doc,
        ["Role", "Responsibility"],
        [
            ("Hieu · program coordination", "Program thesis, venue relationship, learning direction and ecosystem coordination"),
            ("Danh · operations lead", "Run of show, suppliers, crew, venue readiness, participant operations and incident escalation"),
            ("Mentors · to be confirmed", "Product, technical and go-to-market feedback across online and in-person checkpoints"),
            ("Event crew", "Registration, room operations, catering, timekeeping, technical support and documentation"),
            ("DN Blockchain Hub / community collaborators", "Local outreach and community support, subject to final confirmation"),
        ],
        [2760, 6600],
    )
    add_page_break(doc)
    add_section_intro(
        doc,
        "07 / Delivery controls",
        "Safeguards, risks and readiness checks",
        "A short event still needs explicit operating controls. Danh leads delivery, while venue and program responsibilities are confirmed before participants arrive.",
    )
    add_heading(doc, "Participant safeguards", 2)
    safeguards = [
        "Published code of conduct and named escalation contacts.",
        "No overnight venue operations and no expectation of overnight building.",
        "Dietary and accessibility needs collected before catering and room plans are finalized.",
        "Photo consent and project-visibility choices included in registration and submission flows.",
        "First-aid supplies, emergency contacts and clear venue evacuation information.",
    ]
    for item in safeguards:
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Key delivery risks", 2)
    add_data_table(
        doc,
        ["Risk", "Mitigation", "Owner"],
        [
            ("Participants arrive unprepared", "Online readiness gate by 17 September; incomplete teams receive remote support before Friday", "Program"),
            ("Internet instability", "Pre-event load test, venue technical contact and backup connectivity budget", "Ops + venue"),
            ("Lower Sunday attendance", "Submission deadline and team demo slots confirmed Saturday; Sunday ends before lunch", "Ops"),
            ("Demo tracks interfere", "Separate rooms or acoustically separated zones; no competing public-address systems", "Venue"),
            ("Catering variance", "Final attendance confirmation, dietary list and 5% delivery buffer", "Ops"),
            ("No local prize reduces urgency", "Progress checkpoints, public demos, mentor access and clear subtrack continuation value", "Program"),
        ],
        [2100, 5700, 1560],
        font_size=8.5,
    )
    add_heading(doc, "Go / no-go checks", 2)
    checks = [
        ("72 hours before. ", "Confirm venue access, final headcount, dietary list, crew rota, demo tracks and supplier delivery windows."),
        ("Before each opening. ", "Test internet, power, displays, microphones, QR codes, emergency contacts and room circulation."),
        ("Before Sunday demos. ", "Lock track assignments, confirm both mentors, publish the submission deadline and test every project page."),
    ]
    for lead, body in checks:
        add_bullet(doc, lead + body, bullet_num_id, bold_lead=lead)

    # 8. Budget.
    add_page_break(doc)
    add_section_intro(
        doc,
        "08 / Budget",
        "A VND 60 million cash plan, with venue supplied in kind",
        "The budget prioritizes participant food, paid operational responsibility and delivery reliability. It excludes local prizes, venue rental and website-development labor.",
    )
    add_heading(doc, "Catering plan · 80 catered people", 2)
    add_data_table(
        doc,
        ["Item", "Planning basis", "Budget"],
        [
            ("Friday light dinner", "80 × VND 80,000", "VND 6.4m"),
            ("Saturday lunch", "80 × VND 80,000", "VND 6.4m"),
            ("Saturday tea break", "80 × VND 55,000", "VND 4.4m"),
            ("Sunday coffee + light breakfast", "80 × VND 55,000", "VND 4.4m"),
            ("Water and continuous coffee", "Weekend provision", "VND 2.4m"),
            ("Delivery and dietary buffer", "Approximately 5%", "VND 1.6m"),
            ("Catering subtotal", "", "VND 25.6m"),
        ],
        [4050, 2790, 2520],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
    )
    add_heading(doc, "Operations plan", 2)
    add_data_table(
        doc,
        ["Item", "Scope", "Budget"],
        [
            ("Operations lead", "Planning and delivery leadership", "VND 8.0m"),
            ("Event crew", "Four-person support across 2.5 days", "VND 4.5m"),
            ("AV, internet and power backup", "Only for gaps in venue provision", "VND 3.0m"),
            ("Photography and media", "Selected event coverage", "VND 4.0m"),
            ("Printing and signage", "Badges, schedules and partner visibility", "VND 2.5m"),
            ("Supplies, first aid and cleaning buffer", "Event consumables and safety", "VND 2.5m"),
            ("Mentor and local transport", "Reasonable reimbursements", "VND 2.0m"),
            ("Website operating services", "Hosting, email and storage; not development", "VND 1.5m"),
            ("Operations subtotal", "", "VND 28.0m"),
            ("Contingency", "Approximately 10%", "VND 5.4m"),
            ("Planned cash total", "Rounded funding target: VND 60.0m", "VND 59.0m"),
        ],
        [3810, 3030, 2520],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        font_size=8.8,
    )
    add_callout(
        doc,
        "Explicit exclusions",
        "No local prize pool · venue supplied in kind · website development contributed or separately sponsored · mentors unpaid except for reasonable transport · the Superteam Vietnam subtrack prize is external to this event budget.",
        fill=AMBER_FILL,
        accent=AMBER,
    )
    add_heading(doc, "Proposed cash-funding structure", 2)
    add_data_table(
        doc,
        ["Partner type", "Target", "Primary coverage"],
        [
            ("Program / operations partner", "VND 25m", "Operations lead, crew and program delivery"),
            ("Food partner", "VND 25m", "Participant meals, breaks and hydration"),
            ("Two supporting partners", "VND 5m each", "Media, supplies, transport and contingency"),
            ("Venue partner", "In kind", "Space, utilities, agreed AV, security and cleaning"),
        ],
        [3210, 1710, 4440],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    # 9. Measurement and next steps.
    add_page_break(doc)
    add_section_intro(
        doc,
        "09 / Measurement and close",
        "How success will be measured and reported",
        "The organizing team will report outcomes that reflect learning, building and continuation—not vanity attendance alone. The venue partner receives a concise report within seven days of the in-person weekend.",
    )
    add_heading(doc, "Proposed success measures", 2)
    add_data_table(
        doc,
        ["Measure", "Target", "Evidence"],
        [
            ("Builder participation", "70 participants", "Registration and check-in"),
            ("Active teams", "11–12 teams", "Website project workspaces"),
            ("Sunday submissions", "11–12 projects", "Timestamped website submissions"),
            ("Iteration completion", "At least 80% of teams publish two or more meaningful checkpoints", "Project histories"),
            ("Subtrack continuation", "At least 8 teams continue toward the Superteam Vietnam subtrack", "Post-event project updates"),
            ("Partner visibility", "Agreed digital and on-site placements delivered", "Website, event photos and signage record"),
        ],
        [2760, 2520, 4080],
        font_size=8.8,
    )
    add_heading(doc, "Post-event report", 2)
    report_items = [
        "Attendance, team and submission totals.",
        "Project directory with links and short descriptions.",
        "Checkpoint and continuation metrics.",
        "Approved photography and partner-visibility record.",
        "Operational lessons and recommendations for the next program.",
    ]
    for item in report_items:
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Partnership next steps", 2)
    next_steps_num_id = duplicate_numbering_instance(doc, number_num_id)
    next_steps = [
        ("Confirm interest. ", "Agree that the venue can support the dates, capacity and requested operating hours."),
        ("Complete a site and infrastructure check. ", "Validate internet, power, layouts, accessibility, AV and Sunday demo separation."),
        ("Document the in-kind contribution. ", "Record the venue's normal commercial value and the exact services included."),
        ("Exchange brand assets and contacts. ", "Confirm public naming, logo use, technical contact and incident escalation route."),
        ("Approve the operating plan. ", "Sign off the access schedule, catering process, security, cleaning and final room layout."),
        ("Launch public promotion. ", "Publish the venue partnership on the website and participant communications."),
    ]
    for lead, body in next_steps:
        add_numbered(doc, lead + body, next_steps_num_id, bold_lead=lead)
    add_section_intro(
        doc,
        "10 / Confirmation",
        "Venue confirmation and reference notes",
        "A short site review and written record of the in-kind contribution will turn this proposal into an executable venue plan.",
        page_break=True,
    )
    add_callout(
        doc,
        "Requested response",
        "Because the online program opens on 4 September and the in-person weekend begins on 18 September, venue confirmation is requested as soon as possible, ideally by 4 September 2026.",
        fill=GREEN_FILL,
        accent=GREEN,
    )
    add_heading(doc, "Confirm during the site review", 2)
    confirmation_checks = [
        "Legal venue name, address, primary contact and weekend technical contact.",
        "Maximum safe capacity and the agreed team-table layout.",
        "Friday setup access and the exact closing procedure for all three days.",
        "Measured internet capacity, network access method and backup-connectivity permissions.",
        "Power distribution, extension-board rules and cable-management requirements.",
        "Included screens, projectors, microphones, furniture, security and cleaning.",
        "Location and acoustic separation of the two Sunday demo tracks.",
        "Emergency, accessibility, catering-delivery and photography requirements.",
        "Commercial value of the in-kind venue contribution and agreed public recognition.",
    ]
    for item in confirmation_checks:
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Reference notes", 2)
    p = doc.add_paragraph(style="Small Note")
    r = p.add_run("1. Colosseum's published 2026 schedule lists the fall online hackathon as 28 September–2 November 2026: ")
    set_run_font(r, size=9, color=MUTED)
    add_hyperlink(p, "Colosseum Codex — 2026 Hackathons", "https://blog.colosseum.com/2026-hackathons-updraft-course-offline-signer-cli/")
    p = doc.add_paragraph(style="Small Note")
    r = p.add_run("2. Current event tea-break references advertise packages beginning around VND 55,000 per guest; final Da Nang vendor quotations are still required: ")
    set_run_font(r, size=9, color=MUTED)
    add_hyperlink(p, "2026 tea-break pricing reference", "https://circlefood.vn/menu-tea-break/")
    p = doc.add_paragraph(style="Small Note")
    r = p.add_run("3. Da Nang meal and event-photography sources were used only as planning signals; all figures remain provisional until written quotations are obtained: ")
    set_run_font(r, size=9, color=MUTED)
    add_hyperlink(p, "Da Nang meal supplier", "https://www.suatancongnghiepdanang.vn/")
    r = p.add_run(" · ")
    set_run_font(r, size=9, color=MUTED)
    add_hyperlink(p, "Da Nang event photography", "https://danangmedia.net/bao-gia-chup-anh-su-kien-da-nang-2023/")
    add_para(doc, "Superteam Vietnam subtrack eligibility, branding and prize terms remain governed by the applicable program rules and final ecosystem coordination.",
             size=8.8, color=MUTED, italic=True, after=0, line=1.15)

    # Structural metadata and final save.
    doc.settings.element.append(OxmlElement("w:updateFields"))
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_document()
    print(path)
