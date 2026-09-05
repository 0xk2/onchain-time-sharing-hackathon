from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_proposal as bp


ROOT = Path(__file__).parent
SUPERTEAM_OUT = ROOT / "onchain-time-builder-sprint-superteam-vietnam-funding-proposal.docx"
VENUE_OUT = ROOT / "onchain-time-builder-sprint-venue-in-kind-proposal.docx"


def setup_page(doc, audience_label):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False


def setup_document(audience_label, title, subject, keywords):
    doc = Document()
    doc._proposal_audience_label = audience_label
    setup_page(doc, audience_label)
    bp.setup_styles(doc)
    bullet_num_id, number_num_id = bp.configure_numbering(doc)
    core = doc.core_properties
    core.title = title
    core.subject = subject
    core.author = "Onchain Time Builder Sprint organizing team"
    core.keywords = keywords
    core.comments = "Prepared 31 August 2026"
    return doc, bullet_num_id, number_num_id


def add_new_section(doc, kicker, title, lede):
    """Start a clean content page with body-based furniture for stable rendering."""
    masthead = bp.add_para(
        doc,
        f"ONCHAIN TIME BUILDER SPRINT  /  {doc._proposal_audience_label}",
        size=8,
        color=bp.MUTED,
        bold=True,
        font=bp.MONO,
        after=18,
        line=1.0,
        keep=True,
    )
    masthead.paragraph_format.page_break_before = True
    return bp.add_section_intro(doc, kicker, title, lede, page_break=False)


def add_cover(doc, proposal_type, prepared_for, ask_label, ask_value, supporting_rows):
    bp.add_para(
        doc,
        "SUPERTEAM VIETNAM · COLOSSEUM BUILDER PROGRAM",
        size=9,
        color=bp.PURPLE,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        font=bp.MONO,
        after=18,
    )
    bp.add_para(
        doc,
        "Onchain Time\nBuilder Sprint",
        size=30,
        color=bp.INK,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
        line=0.96,
    )
    bp.add_para(
        doc,
        proposal_type,
        size=15,
        color=bp.PURPLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
        line=1.0,
    )
    bp.add_para(
        doc,
        "Learn online. Build together. Iterate toward Colosseum.",
        size=11.5,
        color=bp.MUTED,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=20,
        line=1.1,
    )
    rule = bp.add_para(doc, "", size=1, after=14)
    bp.paragraph_bottom_border(rule, color=bp.PURPLE, size=18, space=8)

    rows = [
        ("PREPARED FOR", prepared_for, "PROGRAM OPENS", "4 September 2026"),
        ("IN-PERSON", "18–20 September 2026", "LOCATION", "Da Nang, Vietnam"),
        ("PARTICIPANTS", "70 builders · 11–12 teams", "FORMAT", "Friday night + 1.5 days"),
        (ask_label, ask_value, supporting_rows[0][0], supporting_rows[0][1]),
    ]
    metadata = doc.add_table(rows=len(rows), cols=2)
    bp.set_table_geometry(metadata, [4680, 4680])
    for row_idx, (l1, v1, l2, v2) in enumerate(rows):
        for col_idx, (label, value) in enumerate(((l1, v1), (l2, v2))):
            cell = metadata.rows[row_idx].cells[col_idx]
            bp.set_cell_shading(cell, bp.SURFACE if row_idx % 2 == 0 else bp.WHITE)
            bp.set_cell_borders(
                cell,
                top={"val": "single", "sz": "4", "color": bp.LINE},
                bottom={"val": "single", "sz": "4", "color": bp.LINE},
                start={"val": "single", "sz": "4", "color": bp.LINE},
                end={"val": "single", "sz": "4", "color": bp.LINE},
            )
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(label + "\n")
            bp.set_run_font(r, name=bp.MONO, size=8, color=bp.PURPLE, bold=True)
            r = p.add_run(value)
            bp.set_run_font(r, size=10.2, color=bp.INK, bold=True)
    bp.add_para(doc, "", size=1, after=10)


def add_program_model(doc, bullet_num_id, funding_audience=False):
    bp.add_heading(doc, "Program model", 2)
    rows = [
        ("Online preparation", "All learning material and build guidance available from 4 September; teams form, define users and begin project pages online."),
        ("In-person acceleration", "Friday kickoff, a focused Saturday build day and a Sunday morning demo-and-feedback session."),
        ("Website workflow", "Learning, team workspaces, checkpoints, feedback and final submission all live on the program website."),
        ("Output", "11–12 working project submissions, each with a defined user, functioning product journey and continuation plan."),
        ("Incentive model", "No local prize. Teams build toward the Superteam Vietnam Colosseum subtrack, with a stated total prize pool of US$10,000."),
    ]
    bp.add_data_table(doc, ["Component", "Plan"], rows, [2400, 6960], font_size=9.0)

    bp.add_heading(doc, "Why this format", 2)
    points = [
        ("Preparation happens before the room. ", "The website carries the learning load, so venue time is used for building, testing and mentor feedback."),
        ("The weekend is deliberately compact. ", "Saturday ends at 17:00 without dinner; Sunday closes at 11:30 without lunch."),
        ("Iteration leads to continuation. ", "Teams are recognized for progress—not ranked locally—and their project histories remain online through the Colosseum cycle."),
    ]
    for lead, body in points:
        bp.add_bullet(doc, lead + body, bullet_num_id, bold_lead=lead)


def add_common_schedule(doc):
    bp.add_heading(doc, "Friday, 18 September · kickoff", 2)
    bp.add_data_table(
        doc,
        ["Time", "Activity", "Result"],
        [
            ("17:30", "Registration and light dinner", "Teams checked in and settled"),
            ("18:15", "Welcome, program purpose and partner acknowledgement", "Shared expectations"),
            ("18:35", "Build prompt, safeguards and website workflow", "Submission requirements understood"),
            ("19:00", "Team plan review and mentor office hours", "Saturday priorities locked"),
            ("20:30", "Checkpoint published on website", "Visible starting point"),
            ("21:00", "Participant close", "No overnight program"),
        ],
        [1500, 4800, 3060],
        font_size=8.7,
    )
    bp.add_heading(doc, "Saturday, 19 September · build and iterate", 2)
    bp.add_data_table(
        doc,
        ["Time", "Activity", "Result"],
        [
            ("09:00", "Stand-up and technical readiness", "Blockers surfaced"),
            ("09:20", "Focused build block", "Core journey implemented"),
            ("11:45", "Checkpoint review", "Evidence and next changes recorded"),
            ("12:00", "Lunch", "On-site meal"),
            ("13:00", "Build, user testing and mentor rotations", "Product iteration"),
            ("15:15", "Tea break", "Short reset"),
            ("15:35", "Final Saturday build block", "Demo path stabilized"),
            ("16:35", "Website checkpoint and Sunday briefing", "Submission plan confirmed"),
            ("17:00", "Close", "No Saturday dinner"),
        ],
        [1500, 4800, 3060],
        font_size=8.7,
    )


def add_demo_model(doc, bullet_num_id):
    bp.add_heading(doc, "Sunday, 20 September · 08:30–11:30", 2)
    bp.add_data_table(
        doc,
        ["Time", "Activity", "Output"],
        [
            ("08:30", "Coffee, light breakfast and final fixes", "Demo-ready project"),
            ("09:15", "Direct website submission deadline", "Stable project page and QR code"),
            ("09:20", "Demo briefing and track assignment", "Teams and audiences positioned"),
            ("09:25–10:55", "Six rounds across two parallel demo tracks", "Capacity for 12 complete demos"),
            ("10:55", "Feedback capture and transition", "Next changes recorded online"),
            ("11:05", "Subtrack continuation plan", "Owners, dates and pathway confirmed"),
            ("11:20", "Partner acknowledgement, photo and close", "Program closes before lunch"),
            ("11:30", "Participants depart", "No Sunday lunch"),
        ],
        [1500, 3900, 3960],
        font_size=8.6,
    )
    bp.add_callout(
        doc,
        "Demo capacity",
        "2 tracks × 6 rounds = 12 team demos. Each 15-minute slot allows approximately 7 minutes for the product, 5 minutes for feedback and 3 minutes for transition.",
        fill=bp.GREEN_FILL,
        accent=bp.GREEN,
    )
    for lead, body in [
        ("Two separated zones. ", "Each track needs a display, power, clear sound and one anchor mentor or facilitator."),
        ("Website-first evidence. ", "Every demo begins from the submitted project page and QR code."),
        ("No local ranking. ", "The close recognizes progress and readiness; the external subtrack governs its own eligibility and prizes."),
    ]:
        bp.add_bullet(doc, lead + body, bullet_num_id, bold_lead=lead)


def build_superteam_proposal():
    doc, bullets, numbers = setup_document(
        "SUPERTEAM VIETNAM FUNDING PROPOSAL",
        "Onchain Time Builder Sprint — Superteam Vietnam Funding Proposal",
        "Program funding and delivery proposal for a two-week online builder program and Da Nang build weekend",
        "Solana, Colosseum, Superteam Vietnam, Da Nang, builder sprint, program funding",
    )
    add_cover(
        doc,
        "Program Funding & Delivery Proposal",
        "Superteam Vietnam",
        "FUNDING REQUEST",
        "VND 60 million",
        [("OPERATIONS LEAD", "Danh")],
    )
    bp.add_callout(
        doc,
        "Decision requested",
        "Approve a VND 60 million program envelope for catering and operations, with the venue sourced separately as an in-kind partnership.",
        fill=bp.GREEN_FILL,
        accent=bp.GREEN,
    )
    bp.add_para(
        doc,
        "Prepared by Hieu and the organizing team · 31 August 2026",
        size=9,
        color=bp.MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
        line=1.1,
    )

    add_new_section(
        doc,
        "01 / Funding case",
        "A compact activation program for the Colosseum pipeline",
        "The Onchain Time Builder Sprint gives Vietnam-based builders two weeks of structured online preparation and a focused Da Nang build weekend. It is designed as part of Superteam Vietnam's Colosseum subtrack effort: teams leave with real products, public iteration evidence and a clear continuation path.",
    )
    bp.add_callout(
        doc,
        "Funding purpose",
        "The requested VND 60 million pays for participant catering and accountable event operations. It does not fund a local prize pool, venue rent or website-development labor.",
    )
    bp.add_heading(doc, "What Superteam Vietnam is funding", 2)
    bp.add_data_table(
        doc,
        ["Funding objective", "Delivery commitment"],
        [
            ("Builder access", "Keep the in-person program free for selected participants and provide essential meals and hydration."),
            ("Operational reliability", "Fund named delivery ownership, event crew, connectivity backup, safety supplies, signage and documentation."),
            ("Subtrack readiness", "Move 11–12 teams from early concepts to submitted working checkpoints and clear next-step plans."),
            ("Ecosystem visibility", "Publish projects, participation evidence, approved media and a post-event outcome report."),
        ],
        [3000, 6360],
    )
    bp.add_heading(doc, "Expected headline outcomes", 2)
    for item in [
        "Approximately 70 participating builders organized into 11–12 teams.",
        "11–12 direct website submissions completed before Sunday demos.",
        "At least 80% of teams publishing two or more meaningful iteration checkpoints.",
        "At least eight teams continuing toward the Superteam Vietnam Colosseum subtrack.",
        "A documented and repeatable online-to-offline builder-program model for future cohorts.",
    ]:
        bp.add_bullet(doc, item, bullets)

    add_new_section(
        doc,
        "02 / Program design",
        "Online learning first; in-person time for execution",
        "The program website is the learning and iteration layer, not merely a registration page. Materials open two weeks before the event, and all projects, checkpoints, feedback and final submissions are recorded directly online.",
    )
    add_program_model(doc, bullets, funding_audience=True)
    bp.add_heading(doc, "Website journey", 2)
    journey_num = bp.duplicate_numbering_instance(doc, numbers)
    for lead, body in [
        ("Learn. ", "Review the prompt, Solana resources, rules and build guidance."),
        ("Form. ", "Create a team, choose a user niche and publish the problem."),
        ("Build. ", "Maintain links, milestones, evidence and an iteration log."),
        ("Submit. ", "Finalize the same project page directly on the website."),
        ("Continue. ", "Use its history and feedback to prepare the subtrack entry."),
    ]:
        bp.add_numbered(doc, lead + body, journey_num, bold_lead=lead)

    add_new_section(
        doc,
        "03 / In-person delivery",
        "One evening, one full day and one focused morning",
        "The schedule protects deep building time while controlling food and operating costs. There is no Saturday dinner, no overnight program and no Sunday lunch.",
    )
    add_common_schedule(doc)

    add_new_section(
        doc,
        "04 / Demo and submission",
        "Twelve teams can demo within two hours",
        "Sunday is a feedback checkpoint, not a local awards ceremony. Two parallel tracks provide enough capacity for every expected team while allowing the event to close at 11:30.",
    )
    add_demo_model(doc, bullets)
    bp.add_heading(doc, "Feedback lenses", 2)
    bp.add_data_table(
        doc,
        ["Lens", "Question"],
        [
            ("User clarity", "Is the target user specific and credible?"),
            ("Value", "Does the time-based exchange solve a real problem?"),
            ("Product", "Can a user complete the essential journey?"),
            ("Evidence", "What did the team learn from real feedback or behavior?"),
            ("Iteration", "What changed across checkpoints, and why?"),
            ("Continuation", "What is the team's next test during Colosseum?"),
        ],
        [2160, 7200],
        font_size=8.7,
    )

    add_new_section(
        doc,
        "05 / Delivery ownership",
        "Clear accountability across program and operations",
        "Danh serves as operations lead, with authority over readiness, suppliers, crew and live-event escalation. Hieu owns program direction, ecosystem coordination and the funding relationship.",
    )
    bp.add_data_table(
        doc,
        ["Role", "Accountability"],
        [
            ("Hieu · program lead", "Program design, Superteam Vietnam coordination, venue relationship, learning direction and final outcome report."),
            ("Danh · operations lead", "Master run of show, procurement, crew roster, venue readiness, participant service and incident escalation."),
            ("Mentors · to be confirmed", "Product, technical and go-to-market feedback online and during the build weekend."),
            ("Four-person event crew", "Registration, room operations, catering, timekeeping, technical support and documentation."),
            ("Venue partner", "In-kind space, utilities, agreed furniture/AV, security, cleaning and venue technical contact."),
            ("Community collaborators", "Local outreach and community support, subject to confirmation."),
        ],
        [2760, 6600],
        font_size=8.8,
    )
    bp.add_heading(doc, "Operating checkpoints", 2)
    for lead, body in [
        ("By 4 September. ", "Confirm funding, venue direction, website readiness and public program message."),
        ("By 11 September. ", "Confirm suppliers, crew, mentors, registration status and draft room plan."),
        ("By 17 September. ", "Complete readiness gate for teams, dietary/accessibility lists, venue test and demo-track plan."),
        ("Each event day. ", "Run internet, power, AV, safety, QR-code and supplier checks before doors open."),
        ("By 27 September. ", "Deliver the outcome report, financial reconciliation and project continuation status."),
    ]:
        bp.add_bullet(doc, lead + body, bullets, bold_lead=lead)

    add_new_section(
        doc,
        "06 / Funding request",
        "VND 60 million for catering and operations",
        "The cash plan is built for 80 catered people: approximately 70 builders plus mentors, crew and partners. Venue rental is excluded because the venue will be requested as an in-kind contribution.",
    )
    bp.add_heading(doc, "Catering · VND 25.6 million", 2)
    bp.add_data_table(
        doc,
        ["Item", "Planning basis", "Budget"],
        [
            ("Friday light dinner", "80 × VND 80,000", "VND 6.4m"),
            ("Saturday lunch", "80 × VND 80,000", "VND 6.4m"),
            ("Saturday tea break", "80 × VND 55,000", "VND 4.4m"),
            ("Sunday coffee + light breakfast", "80 × VND 55,000", "VND 4.4m"),
            ("Water and continuous coffee", "Weekend provision", "VND 2.4m"),
            ("Delivery and dietary buffer", "Approximately 5%", "VND 1.6m"),
            ("Catering subtotal", "", "VND 25.6m"),
        ],
        [4050, 2790, 2520],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
    )
    bp.add_heading(doc, "Operations · VND 28.0 million", 2)
    bp.add_data_table(
        doc,
        ["Item", "Scope", "Budget"],
        [
            ("Operations lead", "Planning and delivery leadership", "VND 8.0m"),
            ("Event crew", "Four-person support across 2.5 days", "VND 4.5m"),
            ("AV, internet and power backup", "Only for gaps in venue provision", "VND 3.0m"),
            ("Photography and media", "Selected event coverage", "VND 4.0m"),
            ("Printing and signage", "Badges, schedules and partner visibility", "VND 2.5m"),
            ("Supplies, first aid and cleaning buffer", "Consumables and safety", "VND 2.5m"),
            ("Mentor and local transport", "Reasonable reimbursements", "VND 2.0m"),
            ("Website operating services", "Hosting, email and storage; not development", "VND 1.5m"),
            ("Operations subtotal", "", "VND 28.0m"),
        ],
        [3810, 3030, 2520],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        font_size=8.7,
    )
    bp.add_heading(doc, "Funding summary", 2)
    bp.add_data_table(
        doc,
        ["Category", "Amount", "Share"],
        [
            ("Catering", "VND 25.6m", "42.7%"),
            ("Operations", "VND 28.0m", "46.7%"),
            ("Contingency and procurement variation", "VND 6.4m", "10.6%"),
            ("Total funding requested", "VND 60.0m", "100%"),
        ],
        [5100, 2400, 1860],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
    )
    bp.add_callout(
        doc,
        "Exclusions",
        "No local prize pool · no venue rent · no website-development labor · mentors unpaid except reasonable transport · the US$10,000 subtrack prize pool is external to this event budget.",
        fill=bp.AMBER_FILL,
        accent=bp.AMBER,
    )

    add_new_section(
        doc,
        "07 / Financial controls",
        "A simple, accountable use-of-funds process",
        "The funding envelope should move fast enough to secure suppliers while remaining easy to reconcile. The following controls are proposed for agreement with Superteam Vietnam.",
    )
    bp.add_heading(doc, "Recommended funding process", 2)
    process_num = bp.duplicate_numbering_instance(doc, numbers)
    for lead, body in [
        ("Approve the envelope. ", "Confirm the VND 60 million cap and eligible categories before supplier commitments."),
        ("Release funds before supplier deadlines. ", "A single pre-event transfer is operationally simplest; any tranche arrangement should still cover deposits and procurement by 7 September."),
        ("Record commitments. ", "Danh maintains the purchase log, supplier quotation/confirmation, owner, due date and payment status."),
        ("Control material changes. ", "Movement above 10% between major categories requires written approval from the program lead and funder contact."),
        ("Reconcile and report. ", "Receipts, invoices, actual-versus-plan variance and unused funds are reported within seven days after the event."),
    ]:
        bp.add_numbered(doc, lead + body, process_num, bold_lead=lead)
    bp.add_heading(doc, "Approval matrix", 2)
    bp.add_data_table(
        doc,
        ["Decision", "Owner", "Evidence"],
        [
            ("Supplier selection and logistics", "Danh", "Quotation or written price confirmation"),
            ("Program-scope and participant decisions", "Hieu", "Program plan and participant records"),
            ("Material category variance", "Hieu + funder contact", "Written approval"),
            ("Payment and expense record", "Danh", "Invoice, receipt or approved exception note"),
            ("Final reconciliation", "Hieu", "Actual-versus-budget statement and report"),
        ],
        [3480, 2160, 3720],
        font_size=8.7,
    )
    bp.add_heading(doc, "Contingency use", 2)
    for item in [
        "Attendance or dietary changes that affect catering quantities.",
        "Internet, display or power backup not included by the final venue.",
        "Last-mile transport, safety or operational requirements discovered during the site check.",
        "Unavoidable supplier variation; contingency is not a discretionary prize or marketing pool.",
    ]:
        bp.add_bullet(doc, item, bullets)

    add_new_section(
        doc,
        "08 / Measurement",
        "Report outcomes that reflect building and continuation",
        "Attendance matters, but the primary evidence is what teams create, how they iterate and whether they continue into the subtrack effort. A concise program and financial report will be delivered within seven days.",
    )
    bp.add_data_table(
        doc,
        ["Measure", "Target", "Evidence"],
        [
            ("Builder participation", "70 participants", "Registration and check-in"),
            ("Active teams", "11–12 teams", "Website team workspaces"),
            ("Sunday submissions", "11–12 projects", "Timestamped website submissions"),
            ("Meaningful iteration", "80% publish 2+ checkpoints", "Project histories"),
            ("Subtrack continuation", "At least 8 teams", "Post-event project updates"),
            ("Delivery quality", "Program closes on schedule with no major incident", "Run-of-show and incident log"),
            ("Financial accountability", "100% of material spend reconciled", "Receipts and variance report"),
        ],
        [2760, 2520, 4080],
        font_size=8.6,
    )
    bp.add_heading(doc, "Post-event report contents", 2)
    for item in [
        "Attendance, team, submission and checkpoint totals.",
        "Project directory with links and concise descriptions.",
        "Continuation status and next milestones for each team.",
        "Approved media and partner-recognition record.",
        "Actual-versus-budget statement with material variance explanations.",
        "Operational lessons and recommendations for a future cohort.",
    ]:
        bp.add_bullet(doc, item, bullets)

    add_new_section(
        doc,
        "09 / Risks and decision",
        "The delivery risks are known and manageable",
        "The main risks are preparation, infrastructure, attendance and cost variance. Each has a specific owner and mitigation path.",
    )
    bp.add_data_table(
        doc,
        ["Risk", "Mitigation", "Owner"],
        [
            ("Teams arrive unprepared", "Online readiness gate by 17 September; targeted remote support before Friday", "Program"),
            ("Venue not confirmed", "Parallel venue outreach, site checklist and decision deadline", "Hieu"),
            ("Internet instability", "Load test, venue technical contact and budgeted backup", "Ops + venue"),
            ("Lower Sunday attendance", "Slots confirmed Saturday; submission deadline tied to demo; close before lunch", "Danh"),
            ("Demo tracks interfere", "Two rooms or acoustically separated zones with independent displays", "Venue"),
            ("Catering variance", "Headcount lock, dietary list and controlled buffer", "Danh"),
            ("No local prize reduces urgency", "Public checkpoints, mentor access and clear subtrack continuation value", "Program"),
        ],
        [2100, 5700, 1560],
        font_size=8.4,
    )
    bp.add_heading(doc, "Requested decisions", 2)
    decision_num = bp.duplicate_numbering_instance(doc, numbers)
    for lead, body in [
        ("Confirm program alignment. ", "Approve the event as part of the Superteam Vietnam Colosseum subtrack effort and confirm the public wording."),
        ("Approve VND 60 million. ", "Authorize the catering and operations envelope and the agreed transfer process."),
        ("Confirm reporting contact. ", "Name the person who approves material variances and receives the final report."),
        ("Confirm subtrack pathway. ", "Validate eligibility, prize-language and direct next steps for participating projects."),
        ("Release program launch. ", "Allow materials and registration communication to go live from 4 September 2026."),
    ]:
        bp.add_numbered(doc, lead + body, decision_num, bold_lead=lead)
    bp.add_callout(
        doc,
        "Timing",
        "Funding and program wording should be confirmed as soon as possible, ideally by 4 September, so suppliers and participant communication can proceed without compression.",
        fill=bp.GREEN_FILL,
        accent=bp.GREEN,
    )
    bp.add_heading(doc, "Reference note", 2)
    p = doc.add_paragraph(style="Small Note")
    r = p.add_run("Colosseum's published 2026 schedule lists the fall online hackathon as 28 September–2 November 2026: ")
    bp.set_run_font(r, size=9, color=bp.MUTED)
    bp.add_hyperlink(p, "Colosseum Codex — 2026 Hackathons", "https://blog.colosseum.com/2026-hackathons-updraft-course-offline-signer-cli/")
    bp.add_para(
        doc,
        "Superteam Vietnam subtrack eligibility, branding and prize terms remain governed by the applicable program rules and final ecosystem coordination.",
        size=8.8,
        color=bp.MUTED,
        italic=True,
        after=0,
        line=1.15,
    )

    doc.settings.element.append(OxmlElement("w:updateFields"))
    doc.save(SUPERTEAM_OUT)
    return SUPERTEAM_OUT


def build_venue_proposal():
    doc, bullets, numbers = setup_document(
        "VENUE IN-KIND PROPOSAL",
        "Onchain Time Builder Sprint — Venue In-Kind Partnership Proposal",
        "In-kind venue partnership proposal for a Da Nang builder sprint",
        "Solana, Superteam Vietnam, Da Nang, builder sprint, venue partnership",
    )
    add_cover(
        doc,
        "Venue In-Kind Partnership Proposal",
        "Prospective Venue Partner",
        "VENUE REQUEST",
        "Space and facilities in kind",
        [("CASH REQUEST", "None")],
    )
    bp.add_callout(
        doc,
        "Partnership request",
        "Provide a safe, accessible and builder-ready venue at no charge for Friday evening, Saturday daytime and Sunday morning. Superteam Vietnam funds the program's catering and operations; no cash sponsorship is requested from the venue partner.",
        fill=bp.GREEN_FILL,
        accent=bp.GREEN,
    )
    bp.add_para(
        doc,
        "Prepared by Hieu and the organizing team · 31 August 2026",
        size=9,
        color=bp.MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
        line=1.1,
    )

    add_new_section(
        doc,
        "01 / The opportunity",
        "Become the physical home of a focused builder program",
        "The Onchain Time Builder Sprint is a two-week online learning and product-iteration program culminating in a compact in-person weekend in Da Nang. It is part of Superteam Vietnam's Colosseum subtrack effort and is designed for building, feedback and continuation—not a local prize ceremony.",
    )
    bp.add_callout(
        doc,
        "Decision requested",
        "Confirm an in-kind venue partnership for up to 80 people, with access from 18–20 September 2026 and the facilities described in this proposal.",
    )
    add_program_model(doc, bullets, funding_audience=False)
    bp.add_heading(doc, "What the venue is—and is not—being asked to provide", 2)
    bp.add_data_table(
        doc,
        ["Requested from venue", "Covered by organizers / Superteam Vietnam"],
        [
            ("Space, utilities, agreed furniture and included AV", "Catering, event staffing and participant operations"),
            ("Weekend access, security, cleaning and venue contact", "Mentors, registration, website and project submission workflow"),
            ("Internet and power suitable for the agreed capacity", "Event supplies, documentation and approved photography"),
            ("Two separated Sunday demo zones", "Program content, timekeeping, feedback and post-event reporting"),
        ],
        [4680, 4680],
        font_size=8.8,
    )

    add_new_section(
        doc,
        "02 / Access and schedule",
        "A bounded request across three calendar days",
        "The program uses Friday evening for kickoff, Saturday for concentrated building and Sunday morning for direct website submission and two parallel demo tracks. There are no overnight activities.",
    )
    bp.add_heading(doc, "Requested venue access", 2)
    bp.add_data_table(
        doc,
        ["Day", "Requested access", "Participant program", "Purpose"],
        [
            ("Fri, 18 Sep", "16:30–21:30", "17:30–21:00", "Setup, registration, kickoff and pack-down"),
            ("Sat, 19 Sep", "08:00–17:30", "09:00–17:00", "Setup, full build day and close; no dinner"),
            ("Sun, 20 Sep", "08:00–12:00", "08:30–11:30", "Breakfast, submissions, demos and close; no lunch"),
        ],
        [1800, 1800, 1800, 3960],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=8.5,
    )
    add_common_schedule(doc)

    add_new_section(
        doc,
        "03 / Space and infrastructure",
        "What a builder-ready venue needs",
        "The ideal venue supports laptop-based team work, stable connectivity and a smooth Sunday transition into two simultaneous demo tracks.",
    )
    bp.add_heading(doc, "Required facilities", 2)
    for lead, body in [
        ("Capacity and layout. ", "Safe capacity for up to 80 people; team tables for laptop work; a central presentation zone; mentor corners; and catering/water space."),
        ("Internet. ", "Stable connectivity for approximately 100 simultaneous devices, a venue technical contact and permission for organizer-provided backup connectivity."),
        ("Power. ", "Sufficient safe outlets for every team, with clear rules for tested extension boards and cable covers."),
        ("Presentation. ", "One central screen or projector with microphone support, plus a display or projector for each Sunday demo track."),
        ("Sunday separation. ", "Two rooms or acoustically separated zones so both 15-minute demo tracks can operate simultaneously."),
        ("Operating environment. ", "Air conditioning, restrooms, drinking-water access, suitable lighting and accessibility information."),
        ("Safety and access. ", "Security, cleaning, entry procedures, emergency contacts, capacity limits and weekend opening/closing responsibility."),
    ]:
        bp.add_bullet(doc, lead + body, bullets, bold_lead=lead)
    bp.add_heading(doc, "Preferred room plan", 2)
    bp.add_data_table(
        doc,
        ["Zone", "Use", "Key requirement"],
        [
            ("Team floor", "11–12 team tables plus circulation", "Power at every table and reliable Wi-Fi"),
            ("Central zone", "Welcome, briefings and Saturday checkpoints", "Main display, microphone and visible clock"),
            ("Demo track A", "Six Sunday team demos", "Display, sound control and 20–30 audience seats"),
            ("Demo track B", "Six Sunday team demos", "Separate display and acoustic separation"),
            ("Mentor corners", "Small-group product and technical support", "Two quieter conversation points"),
            ("Service zone", "Registration, food, water and supplies", "Delivery access without blocking team work"),
        ],
        [2400, 3720, 3240],
        font_size=8.5,
    )

    add_new_section(
        doc,
        "04 / Sunday demos",
        "Two parallel tracks are enough for the expected teams",
        "With 11–12 teams and two hours available, two tracks preserve meaningful product and feedback time while allowing the venue to close before lunch.",
    )
    add_demo_model(doc, bullets)
    bp.add_heading(doc, "Sunday venue reset", 2)
    for item in [
        "Both demo displays, power and network connections tested before 08:30.",
        "Track signs and team order posted before the 09:20 briefing.",
        "No shared public-address audio between tracks during the six demo rounds.",
        "Participant areas cleared by 11:30; organizer pack-down completed by 12:00.",
    ]:
        bp.add_bullet(doc, item, bullets)

    add_new_section(
        doc,
        "05 / Partner value",
        "A visible role in Da Nang's builder pipeline",
        "The venue partner becomes the physical home of a program connecting online learning, local technical talent and a continuation pathway into the Solana ecosystem.",
    )
    bp.add_heading(doc, "Venue partner benefits", 2)
    for lead, body in [
        ("Official recognition. ", "Named as Official Venue Partner on the event website, registration pages and participant communication."),
        ("On-site visibility. ", "Logo placement at registration, the central presentation area and both Sunday demo tracks, subject to brand guidelines."),
        ("Opening participation. ", "A short welcome during Friday kickoff and acknowledgement during Sunday's close."),
        ("Talent access. ", "Direct engagement with approximately 70 builders and visibility into 11–12 emerging project teams."),
        ("Content and media. ", "Access to approved event photographs, project links and a concise post-event impact report."),
        ("Community association. ", "Visible support for practical technology education and product-building activity in Da Nang."),
    ]:
        bp.add_bullet(doc, lead + body, bullets, bold_lead=lead)
    bp.add_heading(doc, "Recognition to confirm", 2)
    bp.add_data_table(
        doc,
        ["Channel", "Proposed recognition"],
        [
            ("Website", "Official Venue Partner name, logo and link"),
            ("Registration and email", "Partner acknowledgement in participant communications"),
            ("On site", "Entrance, main stage and demo-track logo placement"),
            ("Program moments", "Friday welcome and Sunday closing acknowledgement"),
            ("Post-event", "Inclusion in approved photo recap and impact report"),
        ],
        [2700, 6660],
        font_size=8.8,
    )

    add_new_section(
        doc,
        "06 / Delivery and confirmation",
        "A short site review turns the proposal into an operating plan",
        "Danh leads event operations. Hieu manages the venue relationship and program coordination. The venue appoints one primary contact and one weekend technical contact.",
    )
    bp.add_heading(doc, "Participant safeguards", 2)
    for item in [
        "Published code of conduct and named escalation contacts.",
        "No overnight venue operations or expectation of overnight building.",
        "Dietary and accessibility needs collected before final room and catering plans.",
        "Photo consent and project-visibility choices included in registration and submission.",
        "First-aid supplies, emergency contacts and clear evacuation information.",
    ]:
        bp.add_bullet(doc, item, bullets)
    bp.add_heading(doc, "Site confirmation checklist", 2)
    for item in [
        "Legal venue name, address, primary contact and weekend technical contact.",
        "Maximum safe capacity and agreed team-table layout.",
        "Exact access, opening, closing and pack-down procedure for all three days.",
        "Measured internet capacity, network access method and backup permissions.",
        "Power distribution, extension-board rules and cable-management requirements.",
        "Included screens, projectors, microphones, furniture, security and cleaning.",
        "Location and acoustic separation of the two Sunday demo tracks.",
        "Emergency, accessibility, catering-delivery and photography requirements.",
        "Commercial value of the in-kind contribution and agreed public recognition.",
    ]:
        bp.add_bullet(doc, item, bullets)
    bp.add_heading(doc, "Next steps", 2)
    next_num = bp.duplicate_numbering_instance(doc, numbers)
    for lead, body in [
        ("Confirm interest. ", "Agree in principle as soon as possible to support the dates, capacity and operating hours; online materials open on 4 September."),
        ("Complete a site check and room plan. ", "Validate internet, power, layout, access, AV and demo separation, then sign off catering delivery, security, cleaning and event zones."),
        ("Document the contribution. ", "Record the included services and normal commercial value of the venue support."),
        ("Exchange contacts and brand assets. ", "Confirm public naming, logo use and escalation routes."),
    ]:
        bp.add_numbered(doc, lead + body, next_num, bold_lead=lead)

    doc.settings.element.append(OxmlElement("w:updateFields"))
    doc.save(VENUE_OUT)
    return VENUE_OUT


if __name__ == "__main__":
    print(build_superteam_proposal())
    print(build_venue_proposal())
